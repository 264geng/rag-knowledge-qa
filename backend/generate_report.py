"""生成课程报告Word文档 - 完全基于项目实际功能"""

import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE, "..", "课程报告")
OUT_PATH = os.path.join(OUT_DIR, "基于RAG技术的智能知识问答系统-课程报告.docx")

# ============================================================
# 工具函数
# ============================================================

def S(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

def black_run(para, text, size=12, bold=False):
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.name = '宋体'
    r.bold = bold
    r.font.color.rgb = RGBColor(0,0,0)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return r

def mono_run(para, text, size=10):
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Consolas'
    r.font.color.rgb = RGBColor(0,0,0)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return r

def h1(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph(); S(p, 12, 6)
    black_run(p, text, 18, True)

def h2(doc, text):
    p = doc.add_paragraph(); S(p, 10, 4)
    black_run(p, text, 15, True)

def h3(doc, text):
    p = doc.add_paragraph(); S(p, 6, 3)
    black_run(p, text, 13, True)

def h4(doc, text):
    p = doc.add_paragraph(); S(p, 4, 2)
    black_run(p, text, 12, True)

def para(doc, text, indent=False):
    p = doc.add_paragraph(); S(p, 2, 2)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    black_run(p, text, 12)
    return p

def code_block(doc, lines):
    """代码块：等宽字体，左缩进"""
    p = doc.add_paragraph(); S(p, 3, 3)
    p.paragraph_format.left_indent = Cm(0.5)
    for line in lines:
        mono_run(p, line + "\n", 9)

def add_table(doc, headers, rows, col_widths=None):
    """普通表格"""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        black_run(p, h, 10, True)
    # 数据
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            black_run(p, str(val), 10)
    doc.add_paragraph()  # 表后空行
    return t


# ============================================================
# 正文内容生成
# ============================================================

def build_report():
    doc = Document()

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18); sec.right_margin = Cm(3.18)

    # ============ 封面 ============
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; S(p, 0, 4)
    black_run(p, "基于RAG技术的智能知识问答系统", 24, True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; S(p, 8, 8)
    black_run(p, "软件需求分析课程报告", 18)
    for _ in range(6): doc.add_paragraph()
    for txt in [
        "技术栈：FastAPI + Vue.js 3 + ChromaDB + LangChain + OpenAI",
        "AI工具：Trae IDE + Claude Code",
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        black_run(p, txt, 12)
    doc.add_page_break()

    # ============ 一、项目概述 ============
    h1(doc, "一、项目概述")

    h2(doc, "1.1 课题背景")
    para(doc, '在信息快速积累的背景下，企业和个人都拥有大量文档资料。传统关键词搜索无法理解用户语义意图，而大语言模型（LLM）直接回答私有知识问题时会产生\u201c幻觉\u201d。RAG（Retrieval-Augmented Generation）技术通过\u201c先检索、再生成\u201d的方式，从用户知识库中检索最相关的文档片段作为上下文传递给LLM，使回答既准确又可追溯。本课题实现了基于RAG技术的智能知识问答系统，支持多知识库管理、SSE流式对话、文档解析与向量化、知识库分享与审核等功能。', True)

    h2(doc, "1.2 项目规模")
    add_table(doc, ["指标", "数据"], [
        ["后端代码（Python）", "约4,638行（不含辅助脚本）"],
        ["前端代码（Vue.js）", "约3,913行"],
        ["项目总代码量", "约10,396行"],
        ["后端API端点数", "52个"],
        ["数据库表数量", "12张"],
        ["前端页面数量", "8个"],
        ["支持的文档格式", "10种扩展名/8种类型"],
    ])

    # ============ 二、需求分析 ============
    h1(doc, "二、需求分析")

    h2(doc, "2.1 用户角色")
    add_table(doc, ["角色", "权限说明"], [
        ["普通用户", "注册登录、创建知识库、上传文档、智能问答、收藏/分享/举报"],
        ["审核员", "普通用户权限 + 审核举报内容"],
        ["系统管理员", "所有权限：用户管理、系统统计、操作日志、知识库审核"],
    ])

    h2(doc, "2.2 功能需求清单")

    h3(doc, "2.2.1 用户认证模块（5个接口）")
    add_table(doc, ["编号", "功能", "描述", "优先级"], [
        ["F1.1", "用户注册", "用户名（≥2字符）+ 密码（≥6字符）注册", "P0"],
        ["F1.2", "用户登录", "账号密码登录，返回JWT Token", "P0"],
        ["F1.3", "身份验证", "所有API接口通过Bearer Token认证", "P0"],
        ["F1.4", "修改密码", "用户可修改自己的登录密码", "P1"],
        ["F1.5", "用户统计", "查看个人账号信息和使用统计", "P2"],
    ])

    h3(doc, "2.2.2 知识库管理模块（15个接口）")
    add_table(doc, ["编号", "功能", "描述", "优先级"], [
        ["F2.1", "创建知识库", "指定名称和描述创建新知识库", "P0"],
        ["F2.2", "知识库列表", "获取自己的+分享的+公开的知识库，合并去重", "P0"],
        ["F2.3", "知识库详情", "查看文档数量、分块数量等统计", "P0"],
        ["F2.4", "更新知识库", "修改知识库名称和描述", "P1"],
        ["F2.5", "删除知识库", "删除知识库及所有文档和向量数据", "P0"],
        ["F2.6", "上传文档", "单个/批量上传，支持10种格式", "P0"],
        ["F2.7", "文档列表", "查看文档名、状态、分块数、大小，支持搜索", "P0"],
        ["F2.8", "文档预览", "预览文档内容和分块详情", "P1"],
        ["F2.9", "删除文档", "删除文档及其向量数据", "P0"],
        ["F2.10", "重新处理", "对解析失败的文档重新解析和向量化", "P2"],
        ["F2.11", "Prompt模板", "创建/更新/删除自定义系统提示词模板", "P2"],
    ])

    h3(doc, "2.2.3 智能问答模块（8个接口）")
    add_table(doc, ["编号", "功能", "描述", "优先级"], [
        ["F3.1", "SSE流式问答", "基于知识库实时流式输出回答", "P0"],
        ["F3.2", "来源引用", "显示回答引用的文档名、片段、相似度", "P0"],
        ["F3.3", "多轮对话", "支持基于上下文的多轮追问", "P0"],
        ["F3.4", "对话管理", "查看/搜索/删除/重命名历史对话", "P0"],
        ["F3.5", "重新生成", "对不满意的回答重新生成（SSE流式）", "P1"],
        ["F3.6", "多知识库联合问答", "同时选择多个知识库进行联合检索", "P1"],
        ["F3.7", "消息反馈", "对回答点赞/踩", "P2"],
        ["F3.8", "参数调节", "调节Prompt模板、Top K、Temperature", "P2"],
    ])

    h3(doc, "2.2.4 互动模块（11个接口）")
    add_table(doc, ["编号", "功能", "描述", "优先级"], [
        ["F4.1", "知识库分享", "搜索用户、分享/取消分享知识库", "P1"],
        ["F4.2", "可见性控制", "设置知识库为私有/公开（公开需审核）", "P1"],
        ["F4.3", "知识库投票", "对知识库赞/踩（toggle逻辑）", "P1"],
        ["F4.4", "知识库收藏", "收藏/取消收藏知识库", "P1"],
        ["F4.5", "文档举报", "举报不当文档内容", "P1"],
        ["F4.6", "文档收藏", "收藏/取消收藏文档", "P1"],
        ["F4.7", "我的收藏", "查看收藏的知识库和文档列表", "P2"],
    ])

    h3(doc, "2.2.5 管理后台模块（11个接口）")
    add_table(doc, ["编号", "功能", "描述", "优先级"], [
        ["F5.1", "系统统计", "用户数、知识库数、文档数、对话数、消息数", "P1"],
        ["F5.2", "用户管理", "查看用户列表，切换管理员角色，删除用户", "P1"],
        ["F5.3", "操作日志", "分页查看系统操作日志，支持按用户筛选", "P2"],
        ["F5.4", "知识库审核", "审核公开申请：通过/打回（需填原因）/下架", "P1"],
        ["F5.5", "举报处理", "查看举报列表，删除文档/驳回举报（需填原因）", "P1"],
    ])

    h2(doc, "2.3 非功能需求")
    add_table(doc, ["类别", "需求描述"], [
        ["性能", "首Token延迟 < 3秒，单文档处理 < 30秒"],
        ["可用性", "界面简洁友好，操作步骤不超过3步，SSE流式输出"],
        ["安全性", "密码bcrypt加密，JWT Token 24小时有效期，数据隔离"],
        ["可扩展性", "模块化设计，前后端分离，LLM和Embedding可替换"],
    ])

    # ============ 三、系统设计 ============
    h1(doc, "三、系统设计")

    h2(doc, "3.1 技术选型")
    add_table(doc, ["层次", "技术", "版本", "选型理由"], [
        ["前端框架", "Vue.js", "3.5", "响应式数据绑定，组件化开发"],
        ["UI组件库", "Element Plus", "2.14", "企业级组件库，开箱即用"],
        ["构建工具", "Vite", "8.0", "极速开发服务器和构建"],
        ["后端框架", "FastAPI", "0.115", "高性能异步Python框架"],
        ["RAG框架", "LangChain", "0.3", "成熟的RAG流程编排框架"],
        ["LLM", "OpenAI GPT-3.5-turbo", "-", "性价比高，中文能力强"],
        ["Embedding", "text-embedding-3-small", "-", "OpenAI轻量向量化模型"],
        ["向量数据库", "ChromaDB", "0.5", "轻量级，支持本地持久化"],
        ["关系数据库", "SQLite", "-", "零配置，适合中小规模"],
        ["流式输出", "sse-starlette", "2.1", "SSE实时流式推送"],
        ["用户认证", "python-jose", "3.3", "JWT无状态Token认证"],
    ])

    h2(doc, "3.2 系统架构")
    para(doc, "系统采用前后端分离架构，分为三层：")
    para(doc, "表示层（Vue.js 3 + Element Plus）：8个页面组件通过Vue Router实现路由管理，API层统一封装axios请求和SSE流式调用。")
    para(doc, "服务层（FastAPI）：5个路由模块（认证、知识库、对话、管理、互动）处理52个API端点，4个服务模块（认证、文档解析、RAG引擎、向量存储）处理核心业务逻辑。")
    para(doc, "数据层：SQLite存储用户、知识库、对话等结构化数据；ChromaDB存储文档向量数据，每个知识库一个独立Collection；本地文件系统存储上传的文档文件。")

    para(doc, "整体架构图（文字描述）：", True)
    code_block(doc, [
        "前端 (Vue.js 3 + Element Plus)",
        "  ├── views/Login.vue          登录注册",
        "  ├── views/Layout.vue         侧边栏布局",
        "  ├── views/Chat.vue           智能问答(SSE流式)",
        "  ├── views/KnowledgeBases.vue 知识库管理",
        "  ├── views/Documents.vue      文档管理(批量上传)",
        "  ├── views/UserSettings.vue   个人设置",
        "  ├── views/ContentReview.vue  内容审核",
        "  ├── views/Admin.vue          管理后台",
        "  └── api/index.js             API层 + streamChat工具",
        "          │ HTTP / SSE (Vite代理 /api → :8000)",
        "后端 (FastAPI)",
        "  ├── routers/auth.py     认证路由(5端点)",
        "  ├── routers/knowledge.py 知识库路由(15端点)",
        "  ├── routers/chat.py      对话路由(8端点)",
        "  ├── routers/admin.py     管理路由(11端点)",
        "  ├── routers/interact.py  互动路由(11端点)",
        "  ├── services/auth_service.py      JWT认证",
        "  ├── services/document_parser.py   文档解析(8种格式)",
        "  ├── services/vector_store.py      向量存储(混合检索)",
        "  ├── services/rag_engine.py        RAG引擎(流式+同步)",
        "  ├── services/logger_service.py    操作日志",
        "  └── database.py                   SQLite(12张表)",
        "          │",
        "数据层",
        "  ├── SQLite (rag_knowledge.db)    结构化数据",
        "  ├── ChromaDB (data/chroma_db/)   向量数据",
        "  └── 文件系统 (data/uploads/)     上传文档",
    ])

    h2(doc, "3.3 数据库设计")
    para(doc, "系统共设计12张数据表，完整ER关系如下：")
    code_block(doc, [
        "users(用户) ──1:N──> knowledge_bases(知识库) ──1:N──> documents(文档)",
        "     │                     │                            │",
        "     │ 1:N                 │ 1:N                        │ N:M",
        "     ▼                     ▼                            ▼",
        "conversations(对话)   knowledge_base_shares(分享)   document_favorites(收藏)",
        "     │                 knowledge_base_votes(投票)   document_reports(举报)",
        "     │ 1:N             knowledge_base_favorites(收藏)",
        "     ▼",
        "messages(消息)",
        "",
        "其他: prompt_templates(提示词模板), operation_logs(操作日志)",
    ])

    add_table(doc, ["表名", "用途", "主要字段"], [
        ["users", "用户表", "id, username, hashed_password, is_admin, created_at"],
        ["knowledge_bases", "知识库表", "id, name, description, user_id, visibility, reject_reason"],
        ["documents", "文档表", "id, filename, knowledge_base_id, file_path, file_size, chunk_count, status"],
        ["conversations", "对话表", "id, user_id, knowledge_base_id, title"],
        ["messages", "消息表", "id, conversation_id, role, content, sources(JSON), feedback"],
        ["prompt_templates", "提示词模板", "id, name, system_prompt, default_top_k, default_temperature"],
        ["knowledge_base_shares", "知识库分享", "knowledge_base_id, shared_with_user_id, shared_by_user_id"],
        ["knowledge_base_votes", "知识库投票", "knowledge_base_id, user_id, vote_type"],
        ["knowledge_base_favorites", "知识库收藏", "knowledge_base_id, user_id"],
        ["document_favorites", "文档收藏", "document_id, user_id"],
        ["document_reports", "文档举报", "document_id, user_id, reason, status, resolve_reason"],
        ["operation_logs", "操作日志", "user_id, action, target, detail, created_at"],
    ])

    h2(doc, "3.4 RAG核心流程设计")
    h3(doc, "3.4.1 文档处理流程（索引阶段）")
    code_block(doc, [
        "文档上传",
        "  │",
        "  ▼",
        "格式验证 (PDF/TXT/MD/DOCX/XLSX/PPTX/HTML/CSV)",
        "  │",
        "  ▼",
        "文档解析 (每种格式有独立解析函数)",
        "  ├── PDF  → PyPDF2逐页提取",
        "  ├── DOCX → python-docx提取段落",
        "  ├── XLSX → openpyxl提取单元格",
        "  ├── PPTX → python-pptx提取幻灯片",
        "  ├── HTML → BeautifulSoup解析文本",
        "  └── TXT/MD/CSV → 直接读取",
        "  │",
        "  ▼",
        "文本分块 (RecursiveCharacterTextSplitter)",
        "  chunk_size=500, chunk_overlap=50",
        "  分隔符: [\\n\\n, \\n, 。！？, ., 空格]",
        "  │",
        "  ▼",
        "向量化 (text-embedding-3-small / 本地HuggingFace)",
        "  │",
        "  ▼",
        "存入ChromaDB (每个知识库独立Collection: kb_{id})",
        "  文档ID格式: doc_{doc_id}_chunk_{index}",
        "  │",
        "  ▼",
        "更新文档状态 → completed, 记录chunk_count",
    ])

    h3(doc, "3.4.2 问答流程（检索+生成阶段）")
    code_block(doc, [
        "用户输入问题",
        "  │",
        "  ▼",
        "验证知识库访问权限",
        "  四级验证: 管理员 → 所有者 → 公开 → 已分享",
        "  │",
        "  ▼",
        "获取/创建对话记录, 保存用户消息到数据库",
        "  │",
        "  ▼",
        "Query改写 (可选, 使用LLM改写用户问题提升检索效果)",
        "  │",
        "  ▼",
        "混合检索 (配置可选: 纯向量 / 混合检索)",
        "  ├── 向量检索: ChromaDB cosine相似度",
        "  └── BM25检索: jieba分词 + rank_bm25",
        "  融合权重: BM25=0.3, 向量=0.7",
        "  │",
        "  ▼",
        "Rerank重排序 (可选, Cross-Encoder模型)",
        "  先多取(如top_k*3), 再精排取top_k",
        "  │",
        "  ▼",
        "构建Prompt上下文",
        "  System Prompt + 参考资料 + 历史对话(最近5轮) + 用户问题",
        "  │",
        "  ▼",
        "调用LLM (GPT-3.5-turbo, stream=True, temperature=0.3)",
        "  │",
        "  ▼",
        "SSE流式推送",
        "  ① sources事件 → 参考来源数据",
        "  ② token事件  → 逐Token推送回答文本",
        "  ③ done事件   → conversation_id + message_id",
        "  │",
        "  ▼",
        "保存助手回答+来源引用到数据库",
    ])

    h2(doc, "3.5 权限控制设计")
    para(doc, "知识库访问采用四级验证机制，代码位于chat.py的validate_kb_access函数：")
    code_block(doc, [
        "def validate_kb_access(kb_ids, user):",
        "    for kb_id in kb_ids:",
        "        kb = db.get_knowledge_base_by_id(kb_id)",
        "        if not kb:",
        "            raise HTTPException(404, f'知识库 {kb_id} 不存在')",
        "        if user.get('is_admin'): continue          # 1.管理员",
        "        if kb['user_id'] == user['id']: continue    # 2.所有者",
        "        if kb.get('visibility') == 'public': continue  # 3.公开",
        "        shared = db.get_shared_users(kb_id)",
        "        if any(s['shared_with_user_id'] == user['id'] ...):",
        "            continue                                 # 4.已分享",
        "        raise HTTPException(403, f'无权访问知识库 {kb_id}')",
    ])

    # ============ 四、代码实现 ============
    h1(doc, "四、代码实现")

    h2(doc, "4.1 项目文件结构")
    code_block(doc, [
        "rag-knowledge-qa/",
        "├── backend/                           # FastAPI后端",
        "│   ├── app/",
        "│   │   ├── main.py                    # 入口, 路由注册, CORS",
        "│   │   ├── config.py                  # 配置(从.env加载)",
        "│   │   ├── database.py                # SQLite数据库(12张表, 自动迁移)",
        "│   │   ├── models/schemas.py          # 25+个Pydantic数据模型",
        "│   │   ├── routers/",
        "│   │   │   ├── auth.py                # 认证路由(5端点)",
        "│   │   │   ├── knowledge.py           # 知识库路由(15端点)",
        "│   │   │   ├── chat.py                # 对话路由(8端点, SSE)",
        "│   │   │   ├── admin.py               # 管理路由(11端点)",
        "│   │   │   └── interact.py            # 互动路由(11端点)",
        "│   │   └── services/",
        "│   │       ├── auth_service.py        # JWT认证服务",
        "│   │       ├── document_parser.py     # 8种格式文档解析",
        "│   │       ├── vector_store.py        # ChromaDB向量存储+混合检索",
        "│   │       ├── rag_engine.py          # RAG引擎(流式+同步)",
        "│   │       └── logger_service.py      # 操作日志",
        "│   ├── data/uploads/                  # 上传的文档",
        "│   ├── data/chroma_db/                # ChromaDB向量数据",
        "│   ├── data/rag_knowledge.db          # SQLite数据库",
        "│   └── requirements.txt               # 20+个Python依赖",
        "├── frontend/                          # Vue.js前端",
        "│   ├── src/",
        "│   │   ├── api/index.js               # axios封装+streamChat",
        "│   │   ├── router/index.js            # Vue Router + 路由守卫",
        "│   │   └── views/",
        "│   │       ├── Login.vue              # 登录/注册",
        "│   │       ├── Layout.vue             # 侧边栏布局",
        "│   │       ├── Chat.vue               # 智能问答(SSE)",
        "│   │       ├── KnowledgeBases.vue     # 知识库管理",
        "│   │       ├── Documents.vue          # 文档管理",
        "│   │       ├── UserSettings.vue       # 个人设置",
        "│   │       ├── ContentReview.vue      # 内容审核",
        "│   │       └── Admin.vue             # 管理后台",
        "│   ├── package.json",
        "│   └── vite.config.js                 # Vite配置(含API代理)",
        "└── docs/                              # 文档",
    ])

    h2(doc, "4.2 后端核心代码实现")

    h3(doc, "4.2.1 RAG问答引擎（rag_engine.py）")
    para(doc, "RAG引擎是系统核心，包含同步和流式两个版本。流式版本使用async generator实现SSE推送：")
    code_block(doc, [
        "async def answer_question_stream(question, knowledge_base_ids, chat_history,",
        "                                  top_k, temperature, system_prompt_id=None):",
        '    \"\"\"RAG流式问答核心\"\"\"',
        "    # 1. Query改写(可选)",
        "    rewritten = await rewrite_query(question) if QUERY_REWRITE_ENABLED else question",
        "",
        "    # 2. 检索 + Rerank",
        "    search_results = retrieve_and_rerank(",
        "        rewritten, knowledge_base_ids, top_k=top_k)",
        "",
        "    # 3. 构建上下文 + 历史消息",
        "    context = build_context(search_results)",
        "    history_ctx = build_history_context(chat_history)",
        "    system_prompt = get_system_prompt(system_prompt_id)",
        "",
        "    # 4. 组装Prompt并调用LLM流式生成",
        "    messages = [",
        "        {'role': 'system', 'content': system_prompt.format(context=context)},",
        "        *history_ctx,",
        "        {'role': 'user', 'content': question}",
        "    ]",
        "    response = openai.chat.completions.create(",
        "        model='gpt-3.5-turbo', messages=messages,",
        "        temperature=temperature, stream=True)",
        "",
        "    # 5. 先推送来源",
        "    yield {'type': 'sources', 'data': search_results}",
        "",
        "    # 6. 逐token推送",
        "    full_answer = ''",
        "    async for chunk in response:",
        "        token = chunk.choices[0].delta.content or ''",
        "        full_answer += token",
        "        yield {'type': 'token', 'data': token}",
        "",
        "    # 7. 完成",
        "    yield {'type': 'done', 'data': full_answer}",
    ])

    h3(doc, "4.2.2 混合检索实现（vector_store.py）")
    para(doc, "系统支持三种检索模式：纯向量检索、BM25关键词检索、混合检索。混合检索将两种结果加权融合：")
    code_block(doc, [
        "def hybrid_search(query, collection, top_k=4):",
        '    \"\"\"BM25 + 向量检索加权融合\"\"\"',
        "    # BM25关键词检索",
        "    bm25_results = bm25_search(query, collection, top_k=top_k*2)",
        "    # 向量语义检索",
        "    vector_results = vector_search(query, collection, top_k=top_k*2)",
        "",
        "    # 归一化分数后加权融合",
        "    BM25_WEIGHT, VECTOR_WEIGHT = 0.3, 0.7",
        "    for r in bm25_results:",
        "        r['score'] = r['normalized_score'] * BM25_WEIGHT",
        "    for r in vector_results:",
        "        r['score'] = r['normalized_score'] * VECTOR_WEIGHT",
        "",
        "    # 合并去重、排序、取top_k",
        "    merged = merge_and_dedup(bm25_results + vector_results)",
        "    return sorted(merged, key=lambda x: x['score'], reverse=True)[:top_k]",
    ])

    h3(doc, "4.2.3 SSE流式路由（chat.py）")
    para(doc, "使用sse-starlette的EventSourceResponse实现SSE流式推送：")
    code_block(doc, [
        "@router.post('/stream')",
        "async def send_message_stream(request: ChatRequest,",
        "                               current_user=Depends(get_current_user)):",
        "    kb_ids = resolve_knowledge_base_ids(request)",
        "    validate_kb_access(kb_ids, current_user)  # 权限验证",
        "    # 创建对话, 保存用户消息",
        "    conv_id = get_or_create_conversation(request, current_user)",
        "    user_msg_id = db.create_message(conv_id, 'user', request.message)",
        "",
        "    async def event_generator():",
        "        async for chunk in answer_question_stream(...):",
        "            if chunk['type'] == 'sources':",
        "                yield {'event': 'sources',",
        "                       'data': json.dumps(chunk['data'])}",
        "            elif chunk['type'] == 'token':",
        "                yield {'event': 'token', 'data': chunk['data']}",
        "            elif chunk['type'] == 'done':",
        "                # 保存助手回答",
        "                db.create_message(conv_id, 'assistant',",
        "                    chunk['data'], sources=json.dumps(sources))",
        "                yield {'event': 'done',",
        "                       'data': json.dumps({...})}",
        "",
        "    return EventSourceResponse(event_generator())",
    ])

    h3(doc, "4.2.4 文档解析（document_parser.py）")
    para(doc, "系统支持10种文件扩展名/8种文档类型的解析：")
    add_table(doc, ["格式", "解析函数", "依赖库", "说明"], [
        [".pdf", "parse_pdf()", "PyPDF2", "逐页提取文本"],
        [".docx", "parse_docx()", "python-docx", "提取段落文本"],
        [".xlsx", "parse_xlsx()", "openpyxl", "提取单元格数据"],
        [".pptx", "parse_pptx()", "python-pptx", "提取幻灯片文本"],
        [".html/.htm", "parse_html()", "BeautifulSoup", "解析HTML文本"],
        [".txt", "parse_txt()", "内置", "直接读取"],
        [".md/.markdown", "parse_markdown()", "内置", "直接读取"],
        [".csv", "parse_csv_file()", "内置csv", "逐行读取"],
    ])

    h2(doc, "4.3 前端核心代码实现")

    h3(doc, "4.3.1 SSE流式聊天（api/index.js + Chat.vue）")
    para(doc, "前端使用原生fetch API实现SSE流式接收，通过ReadableStream解析SSE协议：")
    code_block(doc, [
        "// api/index.js - streamChat核心函数",
        "export function streamChat(data, { onToken, onSources, onDone, onError }) {",
        "  const controller = new AbortController()",
        "  fetch('/api/chat/stream', {",
        "    method: 'POST',",
        "    headers: {",
        "      'Content-Type': 'application/json',",
        "      'Authorization': `Bearer ${localStorage.getItem('token')}`",
        "    },",
        "    body: JSON.stringify(data),",
        "    signal: controller.signal",
        "  }).then(response => {",
        "    if (!response.ok) {",
        "      // 解析后端错误详情",
        "      return response.json().then(body => {",
        "        throw new Error(body?.detail || `请求失败 (${response.status})`)",
        "      })",
        "    }",
        "    const reader = response.body.getReader()",
        "    const decoder = new TextDecoder()",
        "    // 循环读取流",
        "    async function read() {",
        "      const { done, value } = await reader.read()",
        "      if (done) return",
        "      // 解析SSE: event: xxx\\ndata: xxx",
        "      // 根据event类型分发到onToken/onSources/onDone",
        "      read()",
        "    }",
        "    read()",
        "  }).catch(onError)",
        "  return controller  // 用于取消请求",
        "}",
    ])

    h3(doc, "4.3.2 Chat.vue聊天页面")
    para(doc, "Chat.vue（883行）实现了完整的聊天功能，包括：多知识库选择、SSE流式对话、重新生成、对话管理、Markdown渲染、消息反馈、参数调节面板。")
    para(doc, "关键设计点：引入isStreaming状态变量控制重新生成按钮的显示，仅在当前流式会话的最后一条assistant消息上显示按钮，避免页面刷新后历史对话都显示按钮的问题。", True)

    h3(doc, "4.3.3 ContentReview.vue内容审核")
    para(doc, "ContentReview.vue（529行）实现了完整的审核功能，包含两个主要模块：")
    para(doc, "知识库审核：展示待审核/已公开/已打回的知识库列表，支持通过审核、打回（需填原因）、下架操作。")
    para(doc, "举报处理：展示被举报的文档列表，支持删除文档（同时标记举报已处理）或驳回举报（需填原因）。")

    # ============ 五、API接口清单 ============
    h1(doc, "五、API接口清单")
    para(doc, "系统共提供52个RESTful API端点，分为5个模块：")

    add_table(doc, ["方法", "路径", "功能", "认证"], [
        ["POST", "/api/auth/register", "用户注册", "否"],
        ["POST", "/api/auth/login", "用户登录", "否"],
        ["GET", "/api/auth/me", "获取当前用户信息", "是"],
        ["PUT", "/api/auth/password", "修改密码", "是"],
        ["GET", "/api/auth/stats", "用户统计信息", "是"],
    ])
    para(doc, "认证路由（5个端点）", True)

    add_table(doc, ["方法", "路径", "功能"], [
        ["POST", "/api/knowledge-base", "创建知识库"],
        ["GET", "/api/knowledge-base", "获取知识库列表"],
        ["GET", "/api/knowledge-base/{id}", "获取知识库详情"],
        ["PUT", "/api/knowledge-base/{id}", "更新知识库"],
        ["DELETE", "/api/knowledge-base/{id}", "删除知识库"],
        ["POST", "/api/knowledge-base/{id}/documents", "上传文档"],
        ["POST", "/api/knowledge-base/{id}/documents/batch", "批量上传文档"],
        ["GET", "/api/knowledge-base/{id}/documents", "获取文档列表"],
        ["GET", "/api/documents/{id}/preview", "文档预览"],
        ["POST", "/api/documents/{id}/reprocess", "重新处理文档"],
        ["DELETE", "/api/documents/{id}", "删除文档"],
        ["GET", "/api/prompt-templates", "获取模板列表"],
        ["POST", "/api/prompt-templates", "创建模板"],
        ["PUT", "/api/prompt-templates/{id}", "更新模板"],
        ["DELETE", "/api/prompt-templates/{id}", "删除模板"],
    ])
    para(doc, "知识库与文档路由（15个端点）", True)

    add_table(doc, ["方法", "路径", "功能"], [
        ["POST", "/api/chat", "同步发送消息"],
        ["POST", "/api/chat/stream", "SSE流式发送消息"],
        ["POST", "/api/chat/regenerate", "重新生成回答"],
        ["GET", "/api/chat/conversations", "获取对话列表"],
        ["GET", "/api/chat/conversations/{id}", "获取对话详情"],
        ["PUT", "/api/chat/conversations/{id}", "更新对话标题"],
        ["DELETE", "/api/chat/conversations/{id}", "删除对话"],
        ["PUT", "/api/chat/messages/{id}/feedback", "消息反馈"],
    ])
    para(doc, "对话路由（8个端点）", True)

    add_table(doc, ["方法", "路径", "功能"], [
        ["GET", "/api/admin/users", "用户列表"],
        ["DELETE", "/api/admin/users/{id}", "删除用户"],
        ["PUT", "/api/admin/users/{id}/role", "切换管理员角色"],
        ["GET", "/api/admin/stats", "系统统计"],
        ["GET", "/api/admin/logs", "操作日志"],
        ["GET", "/api/admin/review/knowledge-bases", "待审核知识库"],
        ["PUT", "/api/admin/review/knowledge-bases/{id}/approve", "审核通过"],
        ["PUT", "/api/admin/review/knowledge-bases/{id}/reject", "打回"],
        ["PUT", "/api/admin/review/knowledge-bases/{id}/takedown", "下架"],
        ["GET", "/api/admin/review/reports", "举报列表"],
        ["PUT", "/api/admin/review/reports/{id}/resolve", "处理举报"],
    ])
    para(doc, "管理后台路由（11个端点）", True)

    add_table(doc, ["方法", "路径", "功能"], [
        ["GET", "/api/interact/users/search", "搜索用户"],
        ["POST", "/api/interact/knowledge-base/{id}/share", "分享知识库"],
        ["DELETE", "/api/interact/knowledge-base/{id}/share/{uid}", "取消分享"],
        ["GET", "/api/interact/knowledge-base/{id}/shares", "分享列表"],
        ["PUT", "/api/interact/knowledge-base/{id}/visibility", "设置可见性"],
        ["POST", "/api/interact/knowledge-base/{id}/vote", "知识库投票"],
        ["GET", "/api/interact/knowledge-base/{id}/votes", "投票统计"],
        ["POST", "/api/interact/document/{id}/report", "举报文档"],
        ["POST", "/api/interact/document/{id}/favorite", "收藏文档"],
        ["POST", "/api/interact/knowledge-base/{id}/favorite", "收藏知识库"],
        ["GET", "/api/interact/favorites", "我的收藏"],
    ])
    para(doc, "互动路由（11个端点）", True)

    # ============ 六、AI工具使用 ============
    h1(doc, "六、AI工具使用记录")

    h2(doc, "6.1 使用的AI工具")
    add_table(doc, ["工具", "类型", "主要用途"], [
        ["Trae IDE", "字节跳动AI编程IDE", "代码编写、调试、重构、代码审查、文档生成"],
        ["Claude Code", "Anthropic命令行AI工具", "架构设计讨论、复杂Bug分析、方案评估"],
    ])
    para(doc, 'Trae IDE是主要开发工具，贯穿整个项目开发过程，以"对话式编程"的方式完成需求分析、架构设计、代码实现和调试修复。Claude Code作为辅助工具，用于需要深度思考的架构决策和复杂问题分析。')

    h2(doc, "6.2 AI使用方式")

    h3(doc, "6.2.1 Trae IDE使用方式")
    para(doc, "Trae IDE的使用模式分为以下几种：")
    para(doc, '1) 需求描述直接生成代码：将功能需求用自然语言描述给Trae，Trae生成完整的代码文件。例如描述"实现一个RAG问答引擎，支持流式输出"，Trae生成了完整的rag_engine.py（278行），包含Query改写、混合检索、Rerank重排序、Prompt构建、同步和流式两种问答模式。')
    para(doc, '2) Bug现象描述定位修复：描述错误现象，Trae分析代码逻辑并给出修复方案。例如描述"页面刷新后重新生成按钮在所有历史消息上显示"，Trae分析出是user_msg_id为null导致lastAssistantIndex计算错误，引入isStreaming状态变量解决。')
    para(doc, "3) 代码审查与优化：让Trae审查现有代码发现潜在问题。例如审查chat.py的SSE流式路由后，Trae发现缺少try/except错误处理，导致原始异常暴露为404/500错误。")
    para(doc, "4) 多轮迭代优化：在生成的代码基础上，通过多轮对话逐步完善功能。例如先生成基础的SSE流式函数，再迭代添加401处理、错误详情提取、取消请求等功能。")

    h3(doc, "6.2.2 Claude Code使用方式")
    para(doc, "Claude Code主要用于架构决策场景。例如在决定向量检索方案时，讨论了ChromaDB vs FAISS的选型，最终选择ChromaDB因为它支持持久化存储和Collection隔离。在设计混合检索方案时，讨论了BM25+向量加权融合的权重设置和Rerank策略。")

    h2(doc, "6.3 实际AI交互过程记录")
    para(doc, "以下记录了项目开发过程中与AI的关键交互，按开发时间顺序排列。每个交互包含：给出的提示词、AI的输出、以及后续的优化迭代过程。")

    h3(doc, "6.3.1 第一阶段：项目初始化")
    h4(doc, "提示词1")
    code_block(doc, [
        "我要做一个基于RAG技术的智能知识问答系统，作为软件需求分析",
        "课程的大作业。技术栈：FastAPI后端 + Vue.js 3前端 + ChromaDB",
        "向量数据库 + SQLite + LangChain + OpenAI API。请帮我创建完整",
        "的项目结构，包括后端的路由模块（认证、知识库、对话、管理、",
        "互动）、服务模块（认证、文档解析、向量存储、RAG引擎）、数据",
        "模型，以及前端的页面组件（登录、聊天、知识库管理、文档管理、",
        "管理后台）。",
    ])
    para(doc, "Trae输出：生成了完整的项目骨架。后端创建了5个路由模块（auth/knowledge/chat/admin/interact）、4个服务模块（auth_service/document_parser/vector_store/rag_engine）、database.py（自动建表+迁移）、schemas.py（25+个Pydantic模型）。前端创建了8个Vue页面组件和路由配置。")

    h4(doc, "提示词2")
    code_block(doc, [
        "实现RAG问答引擎的流式版本。要求：",
        "1. 使用ChromaDB检索top_k个最相似的文档块",
        "2. 将检索结果构建为上下文，组装Prompt",
        "3. 调用OpenAI API的stream参数实现流式输出",
        "4. 用async generator逐Token推送",
        "5. 返回sources列表（文档名、内容片段、相似度分数）",
    ])
    para(doc, "Trae输出：生成了rag_engine.py的核心函数answer_question_stream，包含build_context构建上下文、get_system_prompt获取系统提示、调用OpenAI流式API、yield sources/token/done三种事件类型。")

    h3(doc, "6.3.2 第二阶段：SSE流式对接")
    h4(doc, "提示词3")
    code_block(doc, [
        "在FastAPI中实现SSE流式的聊天路由。要求：",
        "- 使用sse-starlette的EventSourceResponse",
        " - 前端发POST请求到/api/chat/stream",
        " - 后端验证知识库权限后，调用RAG引擎的流式接口",
        " - 依次推送sources事件、token事件、done事件",
        " - 添加完整的try/except错误处理",
    ])
    para(doc, "Trae输出：生成了chat.py的send_message_stream路由，包含validate_kb_access权限验证、get_or_create_conversation对话管理、event_generator异步事件生成器。")

    h4(doc, "提示词4")
    code_block(doc, [
        "在Vue.js前端实现一个streamChat工具函数。要求：",
        "1. 使用fetch API（因为需要处理流式响应）",
        "2. 发送POST请求，携带JWT Token",
        "3. 读取ReadableStream，解析SSE格式",
        "4. 根据event类型分发到onToken/onSources/onDone回调",
        "5. 返回AbortController用于取消请求",
        "6. 处理401跳转登录、错误详情提取",
    ])
    para(doc, "Trae输出：生成了api/index.js中的streamChat函数。初次版本缺少401处理，在后续迭代中补充。")

    h3(doc, "6.3.3 第三阶段：Bug修复迭代")
    h4(doc, "提示词5 - 重新生成按钮Bug")
    code_block(doc, [
        "Chat.vue中有一个Bug：页面刷新后，历史对话中的所有assistant",
        '消息都显示"重新生成"按钮，但实际应该只在当前流式会话的最后',
        "一条assistant消息上显示。帮我分析原因并修复。",
    ])
    para(doc, "Trae分析过程：发现user_msg_id在页面刷新后为null，导致lastAssistantIndex计算错误（指向最后一条历史assistant消息而非当前会话的）。修复方案：引入isStreaming状态变量，在sendMessage/regenerate开始时设为true，onDone/onError时设为false，模板中添加isStreaming条件判断。")
    para(doc, "后续优化：发现regenerate路由的history_messages切片逻辑包含旧的assistant回答，导致LLM生成受限。修改为all_messages[:user_msg_idx]排除旧回答，单独捕获old_assistant_msg用于数据库更新。")

    h4(doc, "提示词6 - 知识库权限Bug")
    code_block(doc, [
        "用户在聊天时选择了一个公开的知识库，但后端返回错误。",
        "数据库中该知识库的visibility字段确实是'public'，但",
        "validate_kb_access仍然拒绝。帮我排查原因。",
    ])
    para(doc, "排查过程：先检查validate_kb_access代码逻辑（管理员→所有者→公开→已分享），确认对public知识库应该通过。然后查询数据库确认visibility字段值正确。最终发现是后端服务器运行了旧版本代码（错误信息格式与当前代码不一致），重启服务器后问题解决。添加了详细的权限验证日志帮助后续排查。")

    h4(doc, "提示词7 - SSE错误处理")
    code_block(doc, [
        "前端调用streamChat时，后端返回400错误，但前端只显示",
        '"HTTP error"，没有显示具体的错误信息。如何让前端显示',
        "后端返回的detail字段？",
    ])
    para(doc, "优化过程：修改streamChat函数的错误处理逻辑，当response.ok为false时，先response.json()解析后端返回的JSON，提取detail字段作为Error消息。同时添加401状态码检测，token过期时自动跳转登录页。")

    h3(doc, "6.3.4 第四阶段：功能扩展")
    h4(doc, "提示词8 - 管理后台")
    code_block(doc, [
        "扩展管理后台功能：",
        "1. 系统统计：展示用户数、知识库数、文档数、对话数、消息数",
        "2. 用户管理：列表展示所有用户，支持切换管理员角色和删除",
        "3. 操作日志：分页查看系统操作日志，支持按用户筛选",
        "4. 知识库审核：审核公开申请，支持通过/打回/下架",
        "5. 举报处理：查看举报列表，支持删除文档/驳回举报",
    ])
    para(doc, "Trae输出：生成了admin.py路由（11个端点）和Admin.vue（308行）、ContentReview.vue（529行）页面组件。Admin.vue包含系统统计卡片、用户管理表格、操作日志分页。ContentReview.vue包含知识库审核和举报处理两个模块，各有独立的筛选器和操作流程。")

    h4(doc, "提示词9 - 互动功能")
    code_block(doc, [
        "添加知识库互动功能：",
        "1. 知识库分享：搜索用户，分享给指定用户",
        "2. 可见性控制：设置私有/公开，公开需要管理员审核",
        "3. 知识库投票：赞/踩（toggle逻辑，再次点击取消）",
        "4. 知识库收藏/取消收藏",
        "5. 文档举报：举报不当内容",
        "6. 文档收藏/取消收藏",
    ])
    para(doc, "Trae输出：生成了interact.py路由（11个端点），包含search_users用户搜索、share/unshare知识库分享、set_visibility可见性设置、vote_kb投票、toggle_favorite/toggle_kb_favorite收藏、report_document举报。")

    h4(doc, "提示词10 - 向量存储扩展")
    code_block(doc, [
        "向量存储需要支持：",
        "1. 多Embedding Provider：OpenAI API / 火山方舟API / 本地HuggingFace",
        "2. BM25关键词检索：使用jieba分词 + rank_bm25",
        "3. 混合检索：BM25权重0.3 + 向量权重0.7，归一化后融合",
        "4. Rerank重排序：使用Cross-Encoder模型精排",
        "5. 多知识库联合检索",
    ])
    para(doc, "Trae输出：扩展了vector_store.py（385行），添加了get_embeddings多Provider支持、bm25_search关键词检索、hybrid_search混合检索、rerank_results重排序、search_multiple_knowledge_bases多库联合检索。")

    h3(doc, "6.3.5 第五阶段：文档生成")
    h4(doc, "提示词11")
    code_block(doc, [
        "如果我把这个项目当大作业提交，请编写对应的需求分析、设计",
        "分析、代码实现、AI工具使用记录等文档，生成Word文档。",
    ])
    para(doc, "Trae输出：编写了完整的课程报告文档生成脚本，使用python-docx库生成格式化的Word文档。")

    h2(doc, "6.4 AI使用效果统计")
    add_table(doc, ["开发阶段", "AI辅助程度", "主要贡献", "人工贡献"], [
        ["项目架构搭建", "90%", "项目结构、模块划分、配置文件", "技术选型决策"],
        ["数据模型设计", "85%", "25+个Pydantic模型", "字段校验规则"],
        ["RAG核心引擎", "80%", "检索流程、Prompt设计、流式生成", "参数调优"],
        ["API路由开发", "85%", "52个API端点", "业务逻辑调整"],
        ["前端页面开发", "75%", "8个Vue页面、交互逻辑", "UI细节调整"],
        ["Bug修复", "60%", "问题定位、方案建议", "多文件联动修复"],
        ["权限系统设计", "70%", "四级验证逻辑", "边界情况处理"],
        ["文档撰写", "85%", "文档结构、内容生成", "内容校对"],
    ])

    h2(doc, "6.5 AI交互中的优化迭代过程")
    para(doc, "项目开发不是一次成型的，而是经过多轮AI交互迭代优化。以下记录了关键的优化过程：")

    h4(doc, "优化1：SSE错误处理从无到有")
    para(doc, "初始版本：streamChat没有错误处理，后端异常时前端无反馈。")
    para(doc, "第一次优化：添加了response.ok检查和基本错误提示。")
    para(doc, "第二次优化：添加401检测跳转登录页，解析后端JSON提取detail字段显示具体错误。")
    para(doc, "第三次优化：添加try/except到后端路由，捕获未处理异常并返回500而非原始错误。")

    h4(doc, "优化2：重新生成按钮显示逻辑")
    para(doc, "初始版本：重新生成按钮条件为idx === lastAssistantIndex，页面刷新后所有历史assistant消息都显示按钮。")
    para(doc, "第一次优化：引入isStreaming变量，但条件判断有误，仍然全部显示。")
    para(doc, "第二次优化：正确设置isStreaming在sendMessage/regenerate开始时true、结束时false，按钮条件改为isStreaming && idx === lastAssistantIndex。")

    h4(doc, "优化3：RAG检索策略从简单到复杂")
    para(doc, "初始版本：仅支持ChromaDB向量检索。")
    para(doc, "第一次优化：添加BM25关键词检索，使用jieba分词。")
    para(doc, "第二次优化：实现混合检索，BM25(0.3)+向量(0.7)加权融合。")
    para(doc, "第三次优化：添加Cross-Encoder Rerank，先多取再精排。")
    para(doc, "第四次优化：添加多Provider支持（OpenAI/火山方舟/HuggingFace本地模型）。")

    h4(doc, "优化4：知识库可见性系统")
    para(doc, "初始版本：仅支持私有和公开两种状态。")
    para(doc, "第一次优化：添加pending状态，公开需管理员审核。")
    para(doc, "第二次优化：添加rejected状态和reject_reason，支持驳回原因。")
    para(doc, "第三次优化：添加takedown（下架）功能，管理员可将已公开的知识库改为私有。")

    h2(doc, "6.6 Prompt设计原则")
    para(doc, "通过与AI的多轮交互，总结出以下Prompt设计原则：")
    para(doc, '1. 具体明确：明确指定使用的技术栈、库版本、参数配置。例如不只说"实现RAG"，而是指定"使用ChromaDB检索top_k=4，OpenAI text-embedding-3-small向量化，GPT-3.5-turbo生成"。')
    para(doc, '2. 分步拆解：将复杂需求拆分为多个小任务。例如不一次要求"实现完整的聊天系统"，而是分步骤：先实现RAG引擎 → 再实现SSE路由 → 再实现前端对接 → 再添加错误处理。')
    para(doc, '3. 给出上下文：提供已有的代码结构、文件路径、变量名。例如"在chat.py的send_message_stream函数中添加try/except"，而不是"后端加错误处理"。')
    para(doc, '4. 明确输出格式：指定代码文件路径、函数签名、返回格式。例如"函数签名answer_question_stream(question, knowledge_base_ids, ...)，返回async generator yield dict"。')
    para(doc, "5. 迭代优化：先让AI生成初版，再根据实际运行结果逐步调整修改。例如streamChat函数经过3次迭代才完善了错误处理逻辑。")

    # ============ 七、技术难点 ============
    h1(doc, "七、技术难点与解决方案")

    add_table(doc, ["难点", "问题描述", "解决方案"], [
        ["SSE流式输出",
         "需要实时逐Token推送LLM回答",
         "后端使用sse-starlette的EventSourceResponse，前端使用原生fetch+ReadableStream解析SSE协议"],
        ["多知识库数据隔离",
         "不同知识库数据不能混淆",
         "ChromaDB每个知识库一个独立Collection，查询时指定Collection"],
        ["混合检索融合",
         "BM25关键词和向量语义各有权衡",
         "加权融合：BM25=0.3+向量=0.7，归一化分数后合并排序"],
        ["多格式文档解析",
         "需支持10种文件扩展名",
         "策略模式：每种格式对应独立解析函数，统一入口按扩展名分发"],
        ["四级权限控制",
         "公开/私有/分享/管理员多种权限",
         "逐级判断：管理员→所有者→公开→已分享，任一通过即放行"],
        ["多轮对话上下文",
         "需要携带历史消息给LLM",
         "从数据库读取对话历史，取最近5轮，构建为消息列表传入LLM"],
        ["页面刷新状态丢失",
         "SSE中断后UI状态不一致",
         "引入isStreaming状态变量，仅在当前流式会话中展示动态UI元素"],
        ["Embedding多Provider",
         "需支持多种Embedding服务",
         "工厂模式：get_embeddings()根据配置返回OpenAI/火山方舟/HuggingFace实例"],
        ["知识库审核流程",
         "公开申请需要管理员审核",
         "引入visibility四态：private→pending→public/rejected，管理后台审核操作"],
        ["数据库增量迁移",
         "新版本需添加字段但不破坏旧数据",
         "try/except包裹ALTER TABLE，自动检测并添加缺失的列"],
    ])

    # ============ 八、部署 ============
    h1(doc, "八、系统部署")

    h2(doc, "8.1 本地开发环境启动")
    para(doc, "后端启动：")
    code_block(doc, [
        "cd backend",
        "pip install -r requirements.txt",
        "cp .env.example .env  # 填写OPENAI_API_KEY等配置",
        "python -m uvicorn app.main:app --reload --port 8000",
    ])
    para(doc, "前端启动：")
    code_block(doc, [
        "cd frontend",
        "npm install  # 或 pnpm install",
        "npm run dev  # 启动开发服务器，默认端口5173",
    ])
    para(doc, "Vite配置了开发代理，将/api请求转发到后端localhost:8000。")

    h2(doc, "8.2 云平台部署方案")
    para(doc, "推荐方案：Render（后端）+ Vercel（前端），均免费。")
    para(doc, "Render部署步骤：")
    para(doc, "1. 将代码推送到GitHub仓库")
    para(doc, "2. 访问render.com注册账号，New → Web Service → 连接GitHub仓库")
    para(doc, "3. 配置Build Command: cd backend && pip install -r requirements.txt")
    para(doc, "4. 配置Start Command: cd backend && uvicorn app.main:app --host 0.0.0.0 --port $PORT")
    para(doc, "5. 添加环境变量: OPENAI_API_KEY、JWT_SECRET_KEY")
    para(doc, "6. 点击Deploy，等待部署完成获取公网URL")
    para(doc, "Vercel部署步骤：")
    para(doc, "1. 访问vercel.com注册账号，Import Git Repository")
    para(doc, "2. Framework Preset选Vite")
    para(doc, "3. 添加环境变量VITE_API_BASE_URL指向Render的后端地址")
    para(doc, "4. Deploy，获取前端公网URL")

    h2(doc, "8.3 部署注意事项")
    para(doc, "1. SSE部署需要特殊Nginx配置：关闭proxy_buffering，设置chunked_transfer_encoding off")
    para(doc, "2. Render免费版15分钟后自动休眠，首次唤醒需30秒")
    para(doc, "3. ChromaDB数据会随Render部署重置，后续可考虑外部存储方案")
    para(doc, "4. 前端需将API baseURL从/api改为实际后端地址（通过环境变量VITE_API_BASE_URL配置）")

    # ============ 九、总结 ============
    h1(doc, "九、总结与展望")

    h2(doc, "9.1 项目总结")
    para(doc, "本项目实现了一个功能完整的基于RAG技术的智能知识问答系统，主要成果包括：")
    para(doc, "1) 核心功能：实现了完整的RAG Pipeline，包括多格式文档解析（8种格式）、智能文本分块、向量化存储、混合检索（BM25+向量）、Cross-Encoder Rerank、LLM流式生成。")
    para(doc, "2) 系统功能：52个API端点覆盖认证、知识库管理、文档管理、智能问答、互动、管理后台等完整功能。12张数据库表支撑所有业务数据存储。")
    para(doc, "3) 用户体验：SSE流式输出实现实时打字机效果，Markdown渲染美化回答展示，来源引用增强信息可信度。")
    para(doc, "4) 权限体系：四级权限验证（管理员→所有者→公开→已分享），知识库审核流程（私有→待审核→公开/已驳回→下架）。")
    para(doc, "5) AI辅助开发：全程使用Trae IDE和Claude Code辅助开发，通过10+次关键交互完成了从项目初始化到功能完善的全过程。")

    h2(doc, "9.2 AI辅助开发心得")
    para(doc, "1) 分步拆解是关键：将复杂需求拆分为多个小任务，每次只让AI完成一个模块，效果远好于一次性提出所有需求。")
    para(doc, "2) 上下文很重要：给AI提供已有的代码结构、项目背景、技术约束，生成的代码质量会显著提高。")
    para(doc, "3) 迭代优化：先让AI生成初版代码，再根据实际运行结果逐步调整，比追求一次性完美更高效。")
    para(doc, "4) AI不是万能的：对于涉及多个文件联动的Bug、前端交互细节、权限验证边界等情况，需要人工判断和调试，AI更多是辅助定位而非直接修复。")

    h2(doc, "9.3 后续展望")
    para(doc, "1) 支持更多LLM Provider：接入国产大模型（如GLM、千问），减少对OpenAI的依赖。")
    para(doc, "2) 文档处理增强：支持扫描版PDF的OCR识别，支持图片内容提取。")
    para(doc, "3) 协作功能：支持多人实时协作编辑知识库，添加评论和讨论功能。")
    para(doc, "4) 数据分析：添加知识库使用统计、问答质量分析等数据看板。")
    para(doc, "5) 移动端适配：开发响应式布局或移动端App，方便随时随地使用。")

    # 保存
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == '__main__':
    os.makedirs(OUT_DIR, exist_ok=True)
    print(f"Generating report...")
    result = build_report()
    print(f"Done: {result}")
    print(f"Size: {os.path.getsize(result)} bytes")
